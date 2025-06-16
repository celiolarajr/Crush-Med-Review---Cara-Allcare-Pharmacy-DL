import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from datetime import datetime
import os

def clean_text(text):
    if not isinstance(text, str):
        text = str(text)
    replacements = {
        "–": "-", "—": "-", "’": "'", "‘": "'",
        "“": '"', "”": '"', "é": "e", "á": "a",
        "í": "i", "ó": "o", "ú": "u", "â": "a",
        "ê": "e", "î": "i", "ô": "o", "û": "u",
        "ã": "a", "õ": "o", "ç": "c", "É": "E",
        "Á": "A", "Í": "I", "Ó": "O", "Ú": "U",
        "Â": "A", "Ê": "E", "Î": "I", "Ô": "O",
        "Û": "U", "Ã": "A", "Õ": "O", "Ç": "C"
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

st.set_page_config(
    page_title="Crush Medication Review",
    page_icon="💊",
    layout="wide"
)

@st.cache_data
def load_data():
    if os.path.exists("Crush_Med_Data_Bank_Clean.xlsx"):
        return pd.read_excel("Crush_Med_Data_Bank_Clean.xlsx")
    elif os.path.exists("data/Crush_Med_Data_Bank_Clean.xlsx"):
        return pd.read_excel("data/Crush_Med_Data_Bank_Clean.xlsx")
    else:
        st.warning("Please upload the medication database file")
        st.stop()

def format_date(date_obj):
    if isinstance(date_obj, str):
        try:
            dt = datetime.strptime(date_obj, "%Y-%m-%d")
            return dt.strftime("%d/%m/%Y")
        except:
            return date_obj
    elif isinstance(date_obj, (datetime, )):
        return date_obj.strftime("%d/%m/%Y")
    else:
        try:
            return date_obj.strftime("%d/%m/%Y")
        except:
            return str(date_obj)

def generate_word_report(patient_name, dob, nursing_home, selected_meds, med_data):
    doc = Document()
    
    # LOGO
    if os.path.exists("LogoCaraAllcare.png"):
        doc.add_picture("LogoCaraAllcare.png", width=Inches(2))
    
    # TÍTULO
    doc.add_heading('Medication Crushability Review Report', level=1)
    doc.add_paragraph()
    
    # Patient details (no relatório vem aqui depois do título e logo)
    dob_formatted = format_date(dob)
    report_date = datetime.now().strftime('%d/%m/%Y %H:%M')
    
    doc.add_paragraph(f"Patient: {patient_name}")
    doc.add_paragraph(f"Date of Birth: {dob_formatted}")
    doc.add_paragraph(f"Nursing Home: {nursing_home}")
    doc.add_paragraph(f"Report Date: {report_date}")
    doc.add_paragraph()
    
    # Nota importante
    note = doc.add_paragraph()
    note_run = note.add_run("IMPORTANT: Crushing tablets renders the medication unlicensed. If a liquid formulation is available, this is always the preferred option.")
    note_run.bold = True
    
    doc.add_paragraph()
    
    # TABELA DE MEDICAMENTOS
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Medication', 'Can Be Crushed?', 'Alternative Form', 'Recommendation']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    for med in selected_meds:
        med_info = med_data[med_data['Drug'] == med].iloc[0]
        row_cells = table.add_row().cells
        row_cells[0].text = str(med_info['Drug'])
        row_cells[1].text = str(med_info['Can be Crushed'])
        row_cells[2].text = str(med_info['Alternative form available?'])
        row_cells[3].text = str(med_info['Recommendation'])
    
    doc.add_paragraph()
    
    # Assinaturas
    def add_signature_line(doc, label):
        p = doc.add_paragraph()
        run = p.add_run(label + ": " + "______________________________")
        run.font.size = Pt(12)
    
    add_signature_line(doc, "Form completed by pharmacist")
    add_signature_line(doc, "Nurse signature")
    add_signature_line(doc, "Doctor signature")
    
    doc.add_paragraph()
    doc.add_paragraph(f"Report generated on: {report_date}")
    
    # Rodapé
    section = doc.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Crush Med Review Tool developed by Celio Lara Junior - Pharmacist - PSI 10003245"
    p.alignment = 2  # alinhamento à direita
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    
    return doc

def main():
    st.title("💊 Crush Medication Review Tool")
    
    med_data = load_data()
    
    if "new_meds" not in st.session_state:
        st.session_state["new_meds"] = []
    if "selected_meds" not in st.session_state:
        st.session_state["selected_meds"] = []
    
    # Primeiro: Dados do paciente
    st.header("Patient and Report Details")
    patient_name = st.text_input("Patient Name")
    dob = st.date_input("Date of Birth")
    st.markdown(f"Selected Date: **{dob.strftime('%d/%m/%Y')}**")
    nursing_home = st.text_input("Nursing Home")
    
    st.markdown("---")
    
    # Depois: seleção dos medicamentos
    med_list = med_data['Drug'].dropna().tolist() + [med["Drug"] for med in st.session_state["new_meds"]]
    
    selected = st.multiselect(
        "Select medications for review:",
        options=med_list,
        default=st.session_state["selected_meds"]
    )
    st.session_state["selected_meds"] = selected
    
    st.markdown("---")
    st.header("Add new medication (if not in database)")
    
    with st.form("add_new_med_form", clear_on_submit=True):
        new_med_name = st.text_input("Medication Name")
        new_can_crush = st.selectbox("Can Be Crushed?", options=["Yes", "No", "Unknown"])
        new_alt_form = st.text_input("Alternative Form Available?")
        new_recommendation = st.text_area("Recommendation")
        submitted = st.form_submit_button("Add Medication")
    
    if submitted:
        if not new_med_name.strip():
            st.error("Please enter a medication name.")
        elif new_med_name.strip() in med_list:
            st.warning("Medication already exists in the list.")
        else:
            new_entry = {
                "Drug": new_med_name.strip(),
                "Can be Crushed": new_can_crush,
                "Alternative form available?": new_alt_form.strip(),
                "Recommendation": new_recommendation.strip()
            }
            st.session_state["new_meds"].append(new_entry)
            st.session_state["selected_meds"].append(new_med_name.strip())
            st.success(f"Medication '{new_med_name.strip()}' added.")
                
    st.markdown("---")
    
    if st.button("Generate Word Report"):
        if not st.session_state["selected_meds"]:
            st.error("Please select at least one medication for the report.")
            return
        if not patient_name.strip():
            st.error("Please enter the patient's name.")
            return
        
        med_data_full = pd.concat([med_data, pd.DataFrame(st.session_state["new_meds"])], ignore_index=True)
        med_data_full.drop_duplicates(subset=["Drug"], inplace=True)
        
        word_doc = generate_word_report(patient_name, dob, nursing_home, st.session_state["selected_meds"], med_data_full)
        
        word_buffer = BytesIO()
        word_doc.save(word_buffer)
        word_buffer.seek(0)
        
        st.success("Word report generated!")
        
        st.download_button(
            label="Download Word Report (.docx)",
            data=word_buffer,
            file_name=f"Crush_Report_{patient_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
